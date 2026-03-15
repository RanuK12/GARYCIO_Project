#!/usr/bin/env python3
"""
GARYCIO Proposal Document Generator
Generates DOCX and PDF versions of the budget proposal.
"""

import os
import sys
import subprocess
from pathlib import Path

# Ensure python-docx is available
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

# --- Configuration ---
BASE_DIR = Path(r"C:\Users\emilio\Desktop\Oficina Ranuk\GARYCIO_Project")
DOCS_DIR = BASE_DIR / "docs"
DOCX_OUTPUT = DOCS_DIR / "GARYCIO_Propuesta_Presupuesto.docx"
PDF_OUTPUT = DOCS_DIR / "GARYCIO_Propuesta_Presupuesto.pdf"


def set_cell_shading(cell, color_hex):
    """Set background shading for a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_text(cell, text, bold=False, size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    """Set text in a cell with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Reduce paragraph spacing in cells
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


def add_styled_table(doc, headers, rows, total_row=None, col_widths=None):
    """Add a professionally styled table to the document."""
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows) + (1 if total_row else 0), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    header_color = "1B4F72"  # Dark blue
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_color)
        set_cell_text(cell, header, bold=True, size=10,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER,
                      color=(255, 255, 255))

    # Data rows with alternating colors
    for ri, row_data in enumerate(rows):
        bg = "EBF5FB" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = table.rows[1 + ri].cells[ci]
            set_cell_shading(cell, bg)
            align = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cell, val, size=10, alignment=align)

    # Total row
    if total_row:
        total_color = "D4E6F1"
        row_idx = 1 + len(rows)
        for ci, val in enumerate(total_row):
            cell = table.rows[row_idx].cells[ci]
            set_cell_shading(cell, total_color)
            align = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cell, val, bold=True, size=10, alignment=align)

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Cm(width)

    return table


def add_heading_styled(doc, text, level=1):
    """Add a heading with consistent styling."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(27, 79, 114)  # Dark blue
    return heading


def add_body_text(doc, text, bold=False, italic=False):
    """Add body text with formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text, bold_prefix=""):
    """Add a bullet point."""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Calibri"
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    return p


def generate_docx():
    """Generate the GARYCIO proposal DOCX document."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # --- Title ---
    title = doc.add_heading("GARYCIO", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(27, 79, 114)
        run.font.size = Pt(28)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Propuesta de Presupuesto - Sistema de Gestión para Recolección de Residuos")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(52, 73, 94)
    run.bold = True

    # Client info
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("Cliente: Gary CIO  |  Preparado por: Ranuk  |  Fecha: Marzo 2026")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(127, 140, 141)
    run.italic = True

    doc.add_paragraph()  # Spacer

    # --- Resumen Ejecutivo ---
    add_heading_styled(doc, "Resumen Ejecutivo", level=1)
    add_body_text(doc,
        "El presente documento detalla la propuesta de presupuesto para el desarrollo del sistema "
        "GARYCIO, una plataforma integral de gestión para servicios de recolección de residuos. "
        "El proyecto se divide en dos fases principales, con el objetivo de entregar valor incremental "
        "desde la primera etapa."
    )

    # --- Fase 1 ---
    add_heading_styled(doc, "Fase 1 - MVP (Producto Mínimo Viable)", level=1)
    add_body_text(doc,
        "La primera fase se enfoca en la automatización de comunicaciones y la base operativa del sistema."
    )

    add_heading_styled(doc, "Alcance de Fase 1", level=2)
    add_bullet(doc, "Desarrollo completo del Bot de WhatsApp con flujos conversacionales, "
               "incluyendo sistema de reclamos y avisos automatizados", "Bot WhatsApp: ")
    add_bullet(doc, "Algoritmo de optimización de recorridos para rutas de recolección",
               "Optimización de recorridos: ")
    add_bullet(doc, "Diseño e implementación de base de datos con migración de datos existentes",
               "Base de datos: ")
    add_bullet(doc, "Deploy en servidor de producción y configuración completa",
               "Deploy: ")

    add_heading_styled(doc, "Presupuesto Fase 1", level=2)
    fase1_headers = ["Concepto", "Precio Mercado (USD)", "Precio Amigo (USD)"]
    fase1_rows = [
        ["Bot WhatsApp (desarrollo + flows + reclamos)", "1,200", "700"],
        ["Optimización de recorridos", "500", "300"],
        ["Base de datos + migración de datos", "350", "200"],
        ["Deploy + configuración servidor", "200", "100"],
    ]
    fase1_total = ["TOTAL FASE 1", "2,250", "1,300"]
    add_styled_table(doc, fase1_headers, fase1_rows, total_row=fase1_total,
                     col_widths=[10, 4.5, 4.5])

    doc.add_paragraph()  # Spacer

    # --- Fase 2 ---
    add_heading_styled(doc, "Fase 2 - Plataforma Completa", level=1)
    add_body_text(doc,
        "La segunda fase expande el sistema con herramientas de administración, "
        "aplicación móvil y gestión avanzada de flota."
    )

    add_heading_styled(doc, "Alcance de Fase 2", level=2)
    add_bullet(doc, "Panel web completo para administración y monitoreo en tiempo real",
               "Dashboard web: ")
    add_bullet(doc, "Aplicación móvil nativa para choferes y peones con funcionalidades offline",
               "App móvil: ")
    add_bullet(doc, "Sistema completo de gestión de flota con tracking GPS en tiempo real",
               "Gestión de flota: ")
    add_bullet(doc, "Testing integral y deploy de la plataforma completa",
               "Testing + Deploy: ")

    add_heading_styled(doc, "Presupuesto Fase 2", level=2)
    fase2_headers = ["Concepto", "Precio Mercado (USD)", "Precio Amigo (USD)"]
    fase2_rows = [
        ["Dashboard web administración", "2,500", "1,400"],
        ["App móvil (choferes + peones)", "3,500", "2,000"],
        ["Sistema de gestión de flota + tracking", "1,800", "1,000"],
        ["Testing + Deploy", "500", "300"],
    ]
    fase2_total = ["TOTAL FASE 2", "8,300", "4,700"]
    add_styled_table(doc, fase2_headers, fase2_rows, total_row=fase2_total,
                     col_widths=[10, 4.5, 4.5])

    doc.add_paragraph()  # Spacer

    # --- Resumen Total ---
    add_heading_styled(doc, "Resumen Total del Proyecto", level=1)
    total_headers = ["", "Precio Mercado (USD)", "Precio Amigo (USD)"]
    total_rows = [
        ["Fase 1 - MVP", "2,250", "1,300"],
        ["Fase 2 - Plataforma Completa", "8,300", "4,700"],
    ]
    total_total = ["TOTAL PROYECTO", "10,550", "6,000"]
    add_styled_table(doc, total_headers, total_rows, total_row=total_total,
                     col_widths=[10, 4.5, 4.5])

    # Savings note
    doc.add_paragraph()
    savings = doc.add_paragraph()
    savings.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = savings.add_run("Ahorro total con precio amigo: USD 4,550 (43% de descuento)")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(39, 174, 96)
    run.bold = True

    doc.add_paragraph()  # Spacer

    # --- Costos Recurrentes ---
    add_heading_styled(doc, "Costos Recurrentes (post-lanzamiento)", level=1)
    add_body_text(doc,
        "Una vez desplegado el sistema, se aplican los siguientes costos operativos mensuales:"
    )

    recur_headers = ["Concepto", "Costo"]
    recur_rows = [
        ["Servidor VPS", "USD 15-25/mes"],
        ["Dominio (anual)", "USD 12/año"],
        ["Mantenimiento y soporte", "USD 100/mes (precio amigo)"],
    ]
    add_styled_table(doc, recur_headers, recur_rows, col_widths=[10, 9])

    doc.add_paragraph()  # Spacer

    # --- Forma de Pago ---
    add_heading_styled(doc, "Forma de Pago", level=1)

    add_heading_styled(doc, "Fase 1", level=2)
    add_bullet(doc, "50% al inicio del proyecto")
    add_bullet(doc, "50% al finalizar y entregar Fase 1")

    add_heading_styled(doc, "Fase 2", level=2)
    add_bullet(doc, "Pago 1: al inicio de Fase 2")
    add_bullet(doc, "Pago 2: al alcanzar hito intermedio (entrega parcial)")
    add_bullet(doc, "Pago 3: al finalizar y entregar Fase 2")

    doc.add_paragraph()  # Spacer

    # --- Notas ---
    add_heading_styled(doc, "Notas Importantes", level=1)
    add_bullet(doc, "Los precios no incluyen IVA.")
    add_bullet(doc, "Los plazos de entrega se acordarán al inicio de cada fase.")
    add_bullet(doc, "El precio amigo está sujeto a la relación comercial y confianza mutua.")
    add_bullet(doc, "Cualquier funcionalidad fuera del alcance descrito se cotizará por separado.")
    add_bullet(doc, "Se incluye soporte técnico durante 30 días posteriores a cada entrega sin costo adicional.")

    # --- Footer ---
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Ranuk - Soluciones Tecnológicas")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(127, 140, 141)
    run.italic = True

    # Save
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(DOCX_OUTPUT))
    print(f"DOCX generated: {DOCX_OUTPUT}")
    return doc


def generate_pdf():
    """Generate PDF from the DOCX using available method."""
    print("Attempting PDF generation...")

    # Method 1: Try docx2pdf (requires MS Word)
    try:
        import docx2pdf
        docx2pdf.convert(str(DOCX_OUTPUT), str(PDF_OUTPUT))
        print(f"PDF generated via docx2pdf: {PDF_OUTPUT}")
        return True
    except Exception as e:
        print(f"docx2pdf failed: {e}")

    # Method 2: Try LibreOffice
    try:
        soffice_paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            "soffice",
        ]
        for soffice in soffice_paths:
            try:
                result = subprocess.run(
                    [soffice, "--headless", "--convert-to", "pdf",
                     "--outdir", str(DOCS_DIR), str(DOCX_OUTPUT)],
                    capture_output=True, text=True, timeout=60
                )
                if result.returncode == 0 and PDF_OUTPUT.exists():
                    print(f"PDF generated via LibreOffice: {PDF_OUTPUT}")
                    return True
            except (FileNotFoundError, subprocess.TimeoutExpired):
                continue
        print("LibreOffice not found or failed.")
    except Exception as e:
        print(f"LibreOffice method failed: {e}")

    # Method 3: Try win32com (COM automation with MS Word)
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(DOCX_OUTPUT.resolve()))
        doc.SaveAs(str(PDF_OUTPUT.resolve()), FileFormat=17)  # 17 = wdFormatPDF
        doc.Close()
        word.Quit()
        print(f"PDF generated via MS Word COM: {PDF_OUTPUT}")
        return True
    except Exception as e:
        print(f"win32com/Word COM failed: {e}")

    # Method 4: Generate PDF from scratch with reportlab
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm, mm
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
        )
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "reportlab"])
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm, mm
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
        )
        from reportlab.lib.enums import TA_CENTER, TA_LEFT

    pdf = SimpleDocTemplate(
        str(PDF_OUTPUT),
        pagesize=A4,
        rightMargin=2*cm, leftMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm
    )

    styles = getSampleStyleSheet()
    dark_blue = colors.HexColor("#1B4F72")
    light_blue = colors.HexColor("#EBF5FB")
    total_blue = colors.HexColor("#D4E6F1")
    green = colors.HexColor("#27AE60")

    # Custom styles
    title_style = ParagraphStyle(
        "CustomTitle", parent=styles["Title"],
        fontSize=28, textColor=dark_blue, alignment=TA_CENTER, spaceAfter=6
    )
    subtitle_style = ParagraphStyle(
        "CustomSubtitle", parent=styles["Normal"],
        fontSize=14, textColor=colors.HexColor("#34495E"),
        alignment=TA_CENTER, spaceAfter=4, leading=18
    )
    info_style = ParagraphStyle(
        "Info", parent=styles["Normal"],
        fontSize=10, textColor=colors.HexColor("#7F8C8D"),
        alignment=TA_CENTER, spaceAfter=20
    )
    h1_style = ParagraphStyle(
        "H1", parent=styles["Heading1"],
        fontSize=16, textColor=dark_blue, spaceAfter=10, spaceBefore=16
    )
    h2_style = ParagraphStyle(
        "H2", parent=styles["Heading2"],
        fontSize=13, textColor=dark_blue, spaceAfter=8, spaceBefore=12
    )
    body_style = ParagraphStyle(
        "Body", parent=styles["Normal"],
        fontSize=11, spaceAfter=8, leading=15
    )
    bullet_style = ParagraphStyle(
        "Bullet", parent=styles["Normal"],
        fontSize=11, spaceAfter=4, leftIndent=20, leading=15,
        bulletIndent=10
    )
    savings_style = ParagraphStyle(
        "Savings", parent=styles["Normal"],
        fontSize=12, textColor=green, alignment=TA_CENTER,
        spaceAfter=12, spaceBefore=8
    )
    footer_style = ParagraphStyle(
        "Footer", parent=styles["Normal"],
        fontSize=10, textColor=colors.HexColor("#7F8C8D"),
        alignment=TA_CENTER
    )

    elements = []

    # Title
    elements.append(Paragraph("GARYCIO", title_style))
    elements.append(Paragraph(
        "<b>Propuesta de Presupuesto - Sistema de Gesti\u00f3n para Recolecci\u00f3n de Residuos</b>",
        subtitle_style
    ))
    elements.append(Paragraph(
        "<i>Cliente: Gary CIO  |  Preparado por: Ranuk  |  Fecha: Marzo 2026</i>",
        info_style
    ))
    elements.append(Spacer(1, 12))

    # Resumen Ejecutivo
    elements.append(Paragraph("Resumen Ejecutivo", h1_style))
    elements.append(Paragraph(
        "El presente documento detalla la propuesta de presupuesto para el desarrollo del sistema "
        "GARYCIO, una plataforma integral de gesti\u00f3n para servicios de recolecci\u00f3n de residuos. "
        "El proyecto se divide en dos fases principales, con el objetivo de entregar valor incremental "
        "desde la primera etapa.",
        body_style
    ))

    # Helper function for styled tables
    def make_table(headers, rows, total_row=None, col_widths=None):
        data = [headers] + rows
        if total_row:
            data.append(total_row)

        t = Table(data, colWidths=col_widths or [10*cm, 4*cm, 4*cm])
        style_cmds = [
            ("BACKGROUND", (0, 0), (-1, 0), dark_blue),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("ALIGN", (1, 0), (-1, -1), "CENTER"),
            ("ALIGN", (0, 0), (0, -1), "LEFT"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#BDC3C7")),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ]
        # Alternating row colors
        for i in range(1, len(rows) + 1):
            bg = light_blue if (i - 1) % 2 == 0 else colors.white
            style_cmds.append(("BACKGROUND", (0, i), (-1, i), bg))
        # Total row
        if total_row:
            tr_idx = len(data) - 1
            style_cmds.append(("BACKGROUND", (0, tr_idx), (-1, tr_idx), total_blue))
            style_cmds.append(("FONTNAME", (0, tr_idx), (-1, tr_idx), "Helvetica-Bold"))

        t.setStyle(TableStyle(style_cmds))
        return t

    # Fase 1
    elements.append(Paragraph("Fase 1 - MVP (Producto M\u00ednimo Viable)", h1_style))
    elements.append(Paragraph(
        "La primera fase se enfoca en la automatizaci\u00f3n de comunicaciones y la base operativa del sistema.",
        body_style
    ))
    elements.append(Paragraph("Alcance de Fase 1", h2_style))
    for item in [
        "<b>Bot WhatsApp:</b> Desarrollo completo del Bot de WhatsApp con flujos conversacionales, incluyendo sistema de reclamos y avisos automatizados",
        "<b>Optimizaci\u00f3n de recorridos:</b> Algoritmo de optimizaci\u00f3n de recorridos para rutas de recolecci\u00f3n",
        "<b>Base de datos:</b> Dise\u00f1o e implementaci\u00f3n de base de datos con migraci\u00f3n de datos existentes",
        "<b>Deploy:</b> Deploy en servidor de producci\u00f3n y configuraci\u00f3n completa",
    ]:
        elements.append(Paragraph(f"\u2022 {item}", bullet_style))

    elements.append(Paragraph("Presupuesto Fase 1", h2_style))
    elements.append(make_table(
        ["Concepto", "Precio Mercado (USD)", "Precio Amigo (USD)"],
        [
            ["Bot WhatsApp (desarrollo + flows + reclamos)", "1,200", "700"],
            ["Optimizaci\u00f3n de recorridos", "500", "300"],
            ["Base de datos + migraci\u00f3n de datos", "350", "200"],
            ["Deploy + configuraci\u00f3n servidor", "200", "100"],
        ],
        total_row=["TOTAL FASE 1", "2,250", "1,300"]
    ))
    elements.append(Spacer(1, 12))

    # Fase 2
    elements.append(Paragraph("Fase 2 - Plataforma Completa", h1_style))
    elements.append(Paragraph(
        "La segunda fase expande el sistema con herramientas de administraci\u00f3n, "
        "aplicaci\u00f3n m\u00f3vil y gesti\u00f3n avanzada de flota.",
        body_style
    ))
    elements.append(Paragraph("Alcance de Fase 2", h2_style))
    for item in [
        "<b>Dashboard web:</b> Panel web completo para administraci\u00f3n y monitoreo en tiempo real",
        "<b>App m\u00f3vil:</b> Aplicaci\u00f3n m\u00f3vil nativa para choferes y peones con funcionalidades offline",
        "<b>Gesti\u00f3n de flota:</b> Sistema completo de gesti\u00f3n de flota con tracking GPS en tiempo real",
        "<b>Testing + Deploy:</b> Testing integral y deploy de la plataforma completa",
    ]:
        elements.append(Paragraph(f"\u2022 {item}", bullet_style))

    elements.append(Paragraph("Presupuesto Fase 2", h2_style))
    elements.append(make_table(
        ["Concepto", "Precio Mercado (USD)", "Precio Amigo (USD)"],
        [
            ["Dashboard web administraci\u00f3n", "2,500", "1,400"],
            ["App m\u00f3vil (choferes + peones)", "3,500", "2,000"],
            ["Sistema de gesti\u00f3n de flota + tracking", "1,800", "1,000"],
            ["Testing + Deploy", "500", "300"],
        ],
        total_row=["TOTAL FASE 2", "8,300", "4,700"]
    ))
    elements.append(Spacer(1, 12))

    # Resumen Total
    elements.append(Paragraph("Resumen Total del Proyecto", h1_style))
    elements.append(make_table(
        ["", "Precio Mercado (USD)", "Precio Amigo (USD)"],
        [
            ["Fase 1 - MVP", "2,250", "1,300"],
            ["Fase 2 - Plataforma Completa", "8,300", "4,700"],
        ],
        total_row=["TOTAL PROYECTO", "10,550", "6,000"]
    ))
    elements.append(Spacer(1, 8))
    elements.append(Paragraph(
        "<b>Ahorro total con precio amigo: USD 4,550 (43% de descuento)</b>",
        savings_style
    ))
    elements.append(Spacer(1, 12))

    # Costos Recurrentes
    elements.append(Paragraph("Costos Recurrentes (post-lanzamiento)", h1_style))
    elements.append(Paragraph(
        "Una vez desplegado el sistema, se aplican los siguientes costos operativos mensuales:",
        body_style
    ))
    elements.append(make_table(
        ["Concepto", "Costo"],
        [
            ["Servidor VPS", "USD 15-25/mes"],
            ["Dominio (anual)", "USD 12/a\u00f1o"],
            ["Mantenimiento y soporte", "USD 100/mes (precio amigo)"],
        ],
        col_widths=[10*cm, 8*cm]
    ))
    elements.append(Spacer(1, 12))

    # Forma de Pago
    elements.append(Paragraph("Forma de Pago", h1_style))
    elements.append(Paragraph("Fase 1", h2_style))
    elements.append(Paragraph("\u2022 50% al inicio del proyecto", bullet_style))
    elements.append(Paragraph("\u2022 50% al finalizar y entregar Fase 1", bullet_style))
    elements.append(Paragraph("Fase 2", h2_style))
    elements.append(Paragraph("\u2022 Pago 1: al inicio de Fase 2", bullet_style))
    elements.append(Paragraph("\u2022 Pago 2: al alcanzar hito intermedio (entrega parcial)", bullet_style))
    elements.append(Paragraph("\u2022 Pago 3: al finalizar y entregar Fase 2", bullet_style))
    elements.append(Spacer(1, 12))

    # Notas
    elements.append(Paragraph("Notas Importantes", h1_style))
    for note in [
        "Los precios no incluyen IVA.",
        "Los plazos de entrega se acordar\u00e1n al inicio de cada fase.",
        "El precio amigo est\u00e1 sujeto a la relaci\u00f3n comercial y confianza mutua.",
        "Cualquier funcionalidad fuera del alcance descrito se cotizar\u00e1 por separado.",
        "Se incluye soporte t\u00e9cnico durante 30 d\u00edas posteriores a cada entrega sin costo adicional.",
    ]:
        elements.append(Paragraph(f"\u2022 {note}", bullet_style))

    elements.append(Spacer(1, 20))
    elements.append(Paragraph("<i>Ranuk - Soluciones Tecnol\u00f3gicas</i>", footer_style))

    pdf.build(elements)
    print(f"PDF generated via reportlab: {PDF_OUTPUT}")
    return True


def main():
    print("=" * 60)
    print("GARYCIO Proposal Generator")
    print("=" * 60)

    # Generate DOCX
    print("\n--- Generating DOCX ---")
    generate_docx()

    # Generate PDF
    print("\n--- Generating PDF ---")
    generate_pdf()

    print("\n" + "=" * 60)
    print("Done! Files generated:")
    print(f"  DOCX: {DOCX_OUTPUT}")
    print(f"  PDF:  {PDF_OUTPUT}")
    print("=" * 60)


if __name__ == "__main__":
    main()
